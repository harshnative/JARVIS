# importing docx and os module 
# install docx by pip install 
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
import os


# function for opening the txt file 1
def pathToTxtFile1():
    os.system("cls")

    # as in some case , the draged file does not have "" at the end
    path1 = input("drag and drop the first file here and press enter after that - ")
    if(path1[0] == "\""):

        path1 = path1[1:]
        path1 = path1[:-1]

    # checking if the path is correct
    try:
        fil = open(path1 , "r+")
        fil.close()
    except FileNotFoundError:
        print("oops error loading the file - file is missing :(")
        print("\nTry entering again")
        path1 = pathToTxtFile1()

    # returning the path
    return path1


# function for opening the txt file 1
def pathToTxtFile2():
    os.system("cls")

    # as in some case , the draged file does not have "" at the end
    path2 = input("drag and drop the second file here and press enter after that - ")

    if(path2[0] == "\""):
        path2 = path2[1:]
        path2 = path2[:-1]

    # checking if the path is correct
    try:
        fil = open(path2 , "r+")
        fil.close()
    except FileNotFoundError:
        print("oops error loading the file - file is missing :(")
        print("\nTry entering again")
        path2 = pathToTxtFile2()

    # returning the path
    return path2


# function to perform the compare task
# input 1 in mode in order to compare txt files without case sensitivity i.e HELLO and hello will be same
def readingFile(path1 , path2 , mode):

    totalCompare = 0
    errorCompare = 0
    
    # using docx module
    document = Document()

    #opening the file
    fil = open(path1 , "r+")
    fil1 = open(path2 , "r+")

    #calculating the number of lines in the files
    countfil = 0
    countfil1 = 0

    print("calculating the length of file 1 : ")
    
    for i in fil:
        countfil += 1
    
    os.system("cls")

    print("calculating the length of file 2 : ")
    
    for i in fil1:
        countfil1 += 1

    count = 1
    

    # calculating the length when the two files will be zipped for running for loop
    if(countfil>countfil1):
        totalCompare = countfil
        fromLen = countfil1
    else:
        totalCompare = countfil1
        fromLen = countfil

    # closing and reopening file - as their is some error ocurring if we dont do this - the zipped for loop was not running
    fil.close()
    fil1.close()

    os.system("cls")

    print("starting to compare : ")

    fil = open(path1 , "r+")
    fil1 = open(path2 , "r+")

    # running the for loop as both files zipped so that we can get the lines from both files at the same time 
    for i,j in zip(fil , fil1):
        print("on line {} outOf {}".format(count , fromLen))

        # getting the lines
        string1 = i.strip()
        string2 = j.strip()   
            
        # if the mode is 1 converting both the captuared lines to lower case to remove case sensitivity
        if(mode == 1):
            string1 = string1.lower()
            string2 = string2.lower()
        else:
            pass
            

        # getting the individual words in the lines captured in string 1 and sting 2 - so that we can output the colored words which do not match 
        myList1 = list(string1.split())
        myList2 = list(string2.split())

        myList1Len = len(myList1)
        myList2Len = len(myList2)

        # making the list of same length - inserting " " in the list which has less elements - to avoid any kind of error like while running for loop we may get a[i] - their can be nothing at i position - to remove this error we are doing this 
        if(myList1Len == myList2Len):
            pass

        else:
            if(myList1Len > myList2Len):
                elementsToBeInserted = myList1Len - myList2Len
                for i in range(elementsToBeInserted):
                    myList2.append(" ")
            else:
                elementsToBeInserted = myList2Len - myList1Len
                for i in range(elementsToBeInserted):
                    myList1.append(" ")
        

        # now since both list are of same length and also have same elements so if list1 == list2 - then the string1 will be equal to string2
        if(myList1 == myList2):
            pass
        
        # if the list is not same then looping over each element of both list and comparing individual elements 
        else:
            # using docx
            stringHeading = "Error on line " + str(count) + " : "
            document.add_heading(stringHeading, level=1)
            p = document.add_paragraph(style='List Bullet')
        
            # starting to loop through each element on both list at the same time - as the both list are of same length so index will return same position element from both list
            
            
            # looping over for printing the data of the line in txt1
            for index in range(len(myList1)):
                i = myList1[index]
                j = myList2[index]
                stringToadd = str(i) + " "

                # if the elements match then the output will be in black colour 
                if(i == j):    
                    p.add_run(stringToadd)
                #if donot match then the word will be outputed in red colour
                else:
                    errorCompare = errorCompare + 1
                    run = p.add_run(stringToadd)
                    run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
            
            p = document.add_paragraph(style='List Bullet')

            # looping over again to print the data of line in txt2 
            
            for index in range(len(myList1)):
                i = myList1[index]
                j = myList2[index]

                stringToadd = str(j) + " "

                if(i == j):
                    p.add_run(stringToadd)
                else:
                    errorCompare = errorCompare + 1
                    run = p.add_run(stringToadd)
                    run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)

        # increasing the count to keep tarck were we have reached 
        count += 1
        errorCompare = errorCompare / 2

    trackCount = count - 1
    
    count = 0


    # again closing and opening the file to avoid errors 
    fil.close()
    fil1.close()

    os.system("cls")

    print("outputting extra lines founded : ")

    fil = open(path1 , "r+")
    fil1 = open(path2 , "r+")

    # if the txt1 has more lines than txt2
    if(countfil>countfil1):
        stringHeading = "these are the extra lines after the " + str(countfil1) + " in txt file no1 : "
        document.add_heading(stringHeading, level=1)
        p = document.add_paragraph()
        # looping to get the elements in txt1
        for i in fil:
            # will start printing after the length till we have compared - as we used zip above - we compare only till the length of the smaller txt file 
            if(trackCount <= count):
                string1 = i.strip()
                p.add_run(string1)
                p.add_run("\n")
                errorCompare = errorCompare + 1
            count += 1
    
    # if the txt2 has more lines than txt2
    elif(countfil<countfil1):
        stringHeading = "these are the extra lines after the " + str(countfil) + " in txt file no2 : "
        document.add_heading(stringHeading, level=1)
        p = document.add_paragraph()
        for i in fil1:
            if(trackCount <= count):
                string1 = i.strip()
                p.add_run(string1)
                p.add_run("\n")
                errorCompare = errorCompare + 1
            count += 1
                

    # using docx and finally outputting the data 
    document.save('output_result.docx')
    fil.close()
    fil1.close()

    percentageMatched = 100 - ((errorCompare/totalCompare) * 100)

    return percentageMatched



def mainForTxtCompare():
    print("Welcome , this program is capable of comparing any two TXT files :)")
    print("Lets get started ............ ")
    print("press enter to continue - ")
    input()

    os.system("cls")

    print("0. caseSensitive (hello and HELLO will be different")
    print("1. notCaseSensitive (hello and HELLO will be same")
    mode = int(input("enter the mode for comparing : "))
    path1 = pathToTxtFile1()
    # path1 = "txt1.txt"
    path2 = pathToTxtFile2()
    # path2 = "txt2.txt"
    
    percentageMatched = readingFile(path1 , path2 , mode)

    os.system("cls")

    print("files compared successfully :)")
    print("\n\nfiles matched - {} %".format(percentageMatched))
    print("\n\nthe details are outputed to output_result.docx (you can use any office app to open it)")
    print("\n\npress enter to exit>>>")
    input()

if __name__ == "__main__":
    mainForTxtCompare()